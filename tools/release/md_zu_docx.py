"""Erzeugt ANLEITUNG.docx aus ANLEITUNG.md.

Bewusst schlank gehalten: die Anleitung nutzt nur Überschriften, Absätze,
Aufzählungen, nummerierte Listen, Trennlinien und Inline-Auszeichnung
(**fett**, `code`). Genau das wird hier abgedeckt.

Aufruf aus dem Projektverzeichnis:
    python tools/release/md_zu_docx.py
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

WURZEL = Path(__file__).resolve().parents[2]
QUELLE = WURZEL / "ANLEITUNG.md"
ZIEL = WURZEL / "ANLEITUNG.docx"

# **fett**, *kursiv* und `code` – der Rest bleibt Klartext.
INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")


def schreibe_inline(absatz, text):
    """Zerlegt eine Zeile in Läufe mit passender Auszeichnung."""
    for teil in INLINE.split(text):
        if not teil:
            continue
        if teil.startswith("**") and teil.endswith("**"):
            absatz.add_run(teil[2:-2]).bold = True
        elif teil.startswith("`") and teil.endswith("`"):
            lauf = absatz.add_run(teil[1:-1])
            lauf.font.name = "Consolas"
            lauf.font.size = Pt(10)
        elif teil.startswith("*") and teil.endswith("*"):
            absatz.add_run(teil[1:-1]).italic = True
        else:
            absatz.add_run(teil)


def main():
    if not QUELLE.exists():
        sys.exit(f"Nicht gefunden: {QUELLE}")

    dok = Document()
    for zeile in QUELLE.read_text(encoding="utf-8").splitlines():
        roh = zeile.rstrip()
        text = roh.strip()

        if not text:
            continue

        # Trennlinie
        if text in ("---", "***", "___"):
            p = dok.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("· · ·")
            continue

        # Überschriften
        if text.startswith("#"):
            ebene = len(text) - len(text.lstrip("#"))
            dok.add_heading(text[ebene:].strip(), level=min(ebene, 4))
            continue

        # Aufzählung (eingerückte Punkte werden zur zweiten Ebene)
        if re.match(r"^[-*+] ", text):
            eingerueckt = len(roh) - len(roh.lstrip())
            stil = "List Bullet 2" if eingerueckt >= 2 else "List Bullet"
            schreibe_inline(dok.add_paragraph(style=stil), text[2:])
            continue

        # Nummerierte Liste
        treffer = re.match(r"^\d+\.\s+(.*)$", text)
        if treffer:
            schreibe_inline(dok.add_paragraph(style="List Number"),
                            treffer.group(1))
            continue

        schreibe_inline(dok.add_paragraph(), text)

    dok.save(ZIEL)
    print(f"Geschrieben: {ZIEL}")


if __name__ == "__main__":
    main()
